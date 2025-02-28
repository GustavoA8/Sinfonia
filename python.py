from selenium import webdriver 
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from time import sleep
from docx import Document
from googletrans import Translator

def scrape_youtube_comments(video_url, output_file, target_comments=2000):
    # Configurar o navegador com user-agent
    options = Options()
    options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
    service = Service('C:\\chromedriver-win64\\chromedriver.exe')  # Atualize o caminho do ChromeDriver
    driver = webdriver.Chrome(service=service, options=options)
    driver.get(video_url)

    # Aguardar o carregamento inicial da página
    sleep(5)

    # Configurar variáveis
    actions = ActionChains(driver)
    comments_data = []
    scroll_count = 0
    no_new_comments_count = 0

    while len(comments_data) < target_comments:
        # Rolar a página para baixo
        actions.send_keys(Keys.PAGE_DOWN).perform()
        sleep(5)  # Intervalo maior para garantir o carregamento
        scroll_count += 1

        # Extrair o HTML da página
        soup = BeautifulSoup(driver.page_source, 'html.parser')

        # Adicionar novos comentários e autores à lista
        initial_comment_count = len(comments_data)
        for comment_div, author_div in zip(
            soup.select('ytd-comment-thread-renderer #content-text'),
            soup.select('ytd-comment-thread-renderer #author-text')
        ):
            comment_text = comment_div.text.strip()
            author_name = author_div.text.strip()
            if (author_name, comment_text) not in comments_data:
                comments_data.append((author_name, comment_text))

        # Verificar se novos comentários foram adicionados
        new_comments = len(comments_data) - initial_comment_count

        if new_comments == 0:
            no_new_comments_count += 1
        else:
            no_new_comments_count = 0  # Reiniciar se novos comentários forem encontrados

        # Exibir progresso
        print(f"Progresso: {len(comments_data)} comentários extraídos...")

        # Salvamento incremental a cada 10 comentários
        if len(comments_data) % 10 == 0:
            save_comments_to_word(comments_data, output_file)
            print(f"Salvamento incremental: {len(comments_data)} comentários salvos.")

        # Parar se não houver novos comentários após 10 rolagens consecutivas
        if no_new_comments_count >= 10:
            print("Parada forçada: não há novos comentários após múltiplas rolagens.")
            break

    # Encerrar o navegador
    driver.quit()

    # Salvar os comentários finais
    if len(comments_data) > 0:
        save_comments_to_word(comments_data, output_file)

def save_comments_to_word(comments_data, filename):
    translator = Translator()
    doc = Document()
    doc.add_heading('Comentários do YouTube', level=1)
    
    for i, (author, comment) in enumerate(comments_data, start=1):
        # Detectar o idioma do comentário
        detected_language = translator.detect(comment).lang
        
        # Adicionar o comentário original
        paragraph = doc.add_paragraph()
        paragraph.add_run(f'{i}. {author}: {comment}')
        
        # Traduzir se o idioma não for português
        if detected_language != 'pt':
            try:
                translation = translator.translate(comment, src=detected_language, dest='pt').text
                paragraph.add_run(f' ({translation})').bold = True
            except Exception as e:
                print(f"Erro ao traduzir comentário {i}: {e}")
    
    doc.save(filename)
    print(f"Arquivo salvo como {filename}")

# Parâmetros
VIDEO_URL = "https://www.youtube.com/watch?v=9g406F8sUfU"  # Atualize para o vídeo desejado
OUTPUT_FILE = "comentarios_youtube.docx"

# Chamar a função para extrair comentários
scrape_youtube_comments(VIDEO_URL, OUTPUT_FILE, target_comments=2000)
from docx import Document
from googletrans import Translator

def translate_word_file(input_path, output_path):
    # Carregar o arquivo Word
    doc = Document(input_path)
    translator = Translator()

    # Criar um novo documento para salvar o resultado
    translated_doc = Document()

    # Traduzir todo o conteúdo do documento
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Traduzir o texto do parágrafo
            translated_text = translator.translate(paragraph.text, src='en', dest='pt').text
        else:
            translated_text = ""

        # Adicionar o texto traduzido ao novo documento
        translated_doc.add_paragraph(translated_text)

    # Salvar o documento traduzido
    translated_doc.save(output_path)
    print(f"Tradução concluída! Arquivo salvo em: {output_path}")

# Exemplo de uso
input_file = "SPSRD_2025_Publication.docx"  # Substitua pelo caminho do arquivo original
output_file = "arquivo_traduzido.docx"  # Substitua pelo caminho do arquivo traduzido
translate_word_file(input_file, output_file)

